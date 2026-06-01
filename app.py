import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from PIL import Image
import io
import re

# Initialize Gemini Client
client = genai.Client(api_key=st.secrets["API_KEY"])

# Set wide layout
st.set_page_config(layout="wide", page_title="Stuilder")

# --- PRISTINE INTERFACE STYLING ---
st.markdown("""
    <style>
    /* Global Styles */
    @font-face { font-family: 'Steleton'; src: url('files/steleton.ttf') format('truetype'); }
    html, body, p, label, input, textarea, select, .stMarkdown { font-family: 'Steleton', sans-serif !important; }
    h1, h2, h3, h4, h5, h6 { font-family: 'Steleton', sans-serif !important; text-align: center !important; font-weight: 800 !important; }
    
    /* Layout Helpers */
    .full-center-line { border: 0; height: 1px; background: #31353d; width: 100% !important; margin: 25px 0 !important; }
    .shortcode-badge { background-color: #1e222b; border: 1px solid #ff4b4b; color: #ff4b4b; padding: 2px 6px; border-radius: 4px; font-family: monospace; font-weight: bold; }
    
    /* Center the Qty input and the X button inside their columns */
    div[data-testid="stNumberInput"] { display: flex; justify-content: center; }
    div[data-testid="stButton"] { display: flex; justify-content: center; }
    
    /* Input/Button Styling */
    div[data-testid="stBaseButton-primary"] > button { width: 100% !important; height: 44px !important; background-color: transparent !important; border: 1px solid #ffffff !important; border-radius: 4px !important; font-weight: bold !important; }
    </style>
""", unsafe_allow_html=True)

# Title
st.title("STUILDER")

# --- STATE ---
if "comp_title" not in st.session_state: st.session_state.comp_title = ""
if "components" not in st.session_state: st.session_state.components = [{"id": 0, "quantity": 1, "name": ""}]
if "rulebook_content" not in st.session_state: st.session_state.rulebook_content = ""
if "image_vault" not in st.session_state: st.session_state.image_vault = {}
if "comp_counter" not in st.session_state: st.session_state.comp_counter = 1

# --- SIDEBAR ---
with st.sidebar:
    st.header("Game Info")
    players = st.slider("Max Players", 1, 100, 4)
    game_types = st.multiselect("Game Type", ["Intelligence", "Physical"])
    mechanic_tags = st.multiselect("Mechanic Tags", ["Calculation", "Deduction", "Memory", "Drafting", "Voting", "Manipulation", "Management", "Endurance", "Agility"])
    duration_num = st.number_input("Duration (minutes)", 1, 480, 30)

# --- MAIN UI ---
st.subheader("Game Title")
st.session_state.comp_title = st.text_input("Game Title", value=st.session_state.comp_title, label_visibility="collapsed")
st.markdown("<hr class='full-center-line'>", unsafe_allow_html=True)

st.subheader("Component/s")
with st.container(border=True):
    for idx, comp in enumerate(st.session_state.components):
        c1, c2, c3 = st.columns([1, 11, 0.6])
        with c1: 
            comp["quantity"] = st.number_input("Qty", value=comp["quantity"], key=f"q_{comp['id']}", label_visibility="collapsed")
        with c2: 
            comp["name"] = st.text_input("Name", value=comp["name"], key=f"n_{comp['id']}", label_visibility="collapsed")
        with c3:
            # Centered the X button here
            if st.button("X", key=f"del_{comp['id']}"): st.session_state.components.pop(idx); st.rerun()

if st.button("+ Add Component"):
    st.session_state.components.append({"id": st.session_state.comp_counter, "quantity": 1, "name": ""})
    st.session_state.comp_counter += 1; st.rerun()

st.markdown("<hr class='full-center-line'>", unsafe_allow_html=True)

st.subheader("Mechanic & Gameplay")
st.session_state.rulebook_content = st.text_area("Rules", value=st.session_state.rulebook_content, height=300, label_visibility="collapsed")

st.subheader("Image Asset/s")
uploaded_files = st.file_uploader("Upload", accept_multiple_files=True, label_visibility="collapsed")
if uploaded_files:
    for f in uploaded_files:
        if f.name not in [v['name'] for v in st.session_state.image_vault.values()]:
            key = f"img_{len(st.session_state.image_vault) + 1}"
            st.session_state.image_vault[key] = {"bytes": f.getvalue(), "name": f.name}
    st.rerun()

# Display Vault
for img_code, img_data in st.session_state.image_vault.items():
    st.info(f"Code: {img_code} | {img_data['name']}")

# --- FINAL GENERATION ---
st.markdown("<br>", unsafe_allow_html=True)
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    if st.button("Clean Up", type="primary", use_container_width=True):
        if not st.session_state.rulebook_content.strip():
            st.warning("Please add rules first.")
        else:
            with st.spinner("AI Cleaning & Compiling..."):
                prompt = f"Professional Game Manual Editor. Title: {st.session_state.comp_title}. Text: {st.session_state.rulebook_content}"
                response = client.models.generate_content(model='gemini-1.5-flash', contents=prompt)
                
                doc = Document()
                doc.add_heading(st.session_state.comp_title or "Game Title", 0)
                doc.add_paragraph(f"Players: {players} | Duration: {duration_num} mins")
                doc.add_heading("Components", level=1)
                for c in st.session_state.components:
                    doc.add_paragraph(f"{c['quantity']}x {c['name']}")
                doc.add_heading("Rules", level=1)
                doc.add_paragraph(response.text)
                
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.success("Document Compiled!")
                st.download_button("⬇️ Download Cleaned Word Document", bio, "Final_Game_Book.docx")